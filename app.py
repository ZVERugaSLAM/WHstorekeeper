import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai
import json
import tempfile
import re
import datetime
import io
import os
import time
import mimetypes
import zipfile
from openpyxl.styles import Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

try:
    API_KEY = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=API_KEY)
except FileNotFoundError:
    st.error("Файл secrets.toml не знайдено.")
    st.stop()
except KeyError:
    st.error("Змінна GEMINI_API_KEY не знайдена у secrets.toml.")
    st.stop()

st.set_page_config(page_title="WHO Warehouse OCR", layout="wide", page_icon="📦")

def api_call_with_retry(func, *args, **kwargs):
    max_retries = 3
    for attempt in range(max_retries):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            error_str = str(e)
            if "429" in error_str or "Quota exceeded" in error_str:
                if attempt < max_retries - 1:
                    match = re.search(r'retry_delay\s*\{\s*seconds:\s*(\d+)\s*\}', error_str)
                    wait_time = int(match.group(1)) + 5 if match else 20
                    st.toast(f"⏳ Ліміт API. Очікування {wait_time} сек... (Спроба {attempt + 1}/{max_retries})")
                    time.sleep(wait_time)
                else:
                    raise e
            else:
                raise e

def process_document_with_gemini(file_path):
    try:
        uploaded_doc = api_call_with_retry(genai.upload_file, path=file_path)
        model = genai.GenerativeModel(model_name="gemini-2.5-flash")
        prompt = """
        You are a logistics assistant at WHO. Analyze the scanned warehouse documents.
        Extract the general document information and ALL listed items.
        Return STRICTLY in VALID JSON format.
        If information is missing, leave the value empty ("").

        Rules for specific fields:
        - "po_number": Extract ONLY the numeric part of the Purchase Order.
        - "exp_date": Format strictly as DD.MM.YYYY.

        Format the output as a valid JSON object with this structure:
        {
            "act_number": "Extract if present",
            "po_number": "Numeric part only",
            "supplier_name": "Supplier or vendor name",
            "invoice_info": "Invoice/Waybill number and date",
            "number_of_parcels": "Total number of parcels/pallets/boxes",
            "items": [
                {
                    "item_name_eng": "FULL English description",
                    "item_name_ukr": "Accurate Ukrainian translation",
                    "who_item_code": "Extract if present",
                    "who_catalogue_item_name": "Extract if present",
                    "batch": "Batch/Lot number",
                    "exp_date": "Strictly DD.MM.YYYY",
                    "quantity": "Numeric quantity"
                }
            ]
        }
        """
        response = api_call_with_retry(model.generate_content, [prompt, uploaded_doc])
        json_text = response.text
        json_match = re.search(r'```json\n(.*?)\n```', json_text, re.DOTALL)
        if json_match:
            json_text = json_match.group(1)
        data = json.loads(json_text)
        
        # Ініціалізація пустих полів
        manual_fields = ["or_number", "project", "task", "award", "award_end_date", "donor", "requester", "wh"]
        for field in manual_fields:
            if field not in data:
                data[field] = ""
                
        if "items" not in data or not isinstance(data["items"], list):
            data["items"] = []
            
        return data
    except Exception as e:
        st.error(f"OCR Error: {e}")
        return {}

def process_packing_lists(file_paths):
    try:
        model = genai.GenerativeModel(model_name="gemini-2.5-flash")
        prompt = """
        You are a logistics data extraction assistant. Analyze packing list images representing multiple boxes on a pallet.
        Extract the general module information and ALL listed items across all provided pages.
        Always translate the English item description to Ukrainian for the 'item_description_ukr' field.
        If any data is missing on the document, leave the value empty ("").
        Dates must be strictly in DD.MM.YYYY format.

        Format the output STRICTLY as a valid JSON object with the following structure:
        {
          "module_name": "Extract the overall module or kit name (e.g., IEHK 2017, NCDK 2022 MODULE MEDICINES)",
          "module_batch": "Extract the overall module batch number (e.g., MPL00010680)",
          "items": [
            {
              "carton_no": "Carton or box number",
              "item_no": "Item code or number",
              "quantity": "Numeric quantity",
              "packing_unit": "UOM or packing unit (e.g., pack, vial, kit)",
              "item_description_ukr": "Accurate Ukrainian translation of the item description",
              "item_description_eng": "Full English item description",
              "batch_no": "Item batch number",
              "man_date": "Manufacturing date",
              "exp_date": "Expiry date"
            }
          ]
        }
        """
        content_request = [prompt]
        for fp in file_paths:
            mime_type, _ = mimetypes.guess_type(fp)
            if mime_type == 'application/pdf':
                uploaded_doc = api_call_with_retry(genai.upload_file, path=fp)
                content_request.append(uploaded_doc)
                time.sleep(2)
            else:
                with open(fp, "rb") as f:
                    doc_data = f.read()
                content_request.append({"mime_type": mime_type or "image/jpeg", "data": doc_data})
        response = api_call_with_retry(model.generate_content, content_request)
        json_text = response.text
        json_match = re.search(r'```json\n(.*?)\n```', json_text, re.DOTALL)
        if json_match:
            json_text = json_match.group(1)
        data = json.loads(json_text)
        return data
    except Exception as e:
        st.error(f"Packing List OCR Error: {e}")
        return {}

def set_cell_text(cell, text, bold=False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold

def get_total_quantity(items):
    try:
        total = sum(float(str(i.get('Quantity', '0')).replace(',', '.').replace(' ', '')) for i in items if str(i.get('Quantity', '')).strip())
        return str(int(total)) if total.is_integer() else str(total)
    except Exception:
        return "See attached list"

def generate_files_in_memory(data):
    base_name = f"ACT_{data.get('act_number', 'XXX')}_PO_{data.get('po_number', 'XXX')}".replace("/", "-")
    
    # Формування рядків для Excel
    excel_rows = []
    for item in data.get('items', []):
        excel_rows.append({
            "ACT": data.get('act_number'), "PO": data.get('po_number'), "OR": data.get('or_number'),
            "Item Name [UKR]": item.get('Item Name [UKR]', ''), 
            "Item Name [ENG]": item.get('Item Name [ENG]', ''),
            "WHO Item code": item.get('WHO Item code', ''), 
            "WHO Catalogue Item Name": item.get('WHO Catalogue Item Name', ''),
            "Batch": item.get('Batch', ''), 
            "Exp.date": item.get('Exp.date', ''), 
            "Quantity": item.get('Quantity', ''), 
            "Project": data.get('project'), "Task": data.get('task'), "Award": data.get('award'), 
            "Award end date": data.get('award_end_date'), "Donor": data.get('donor'), 
            "Requester": data.get('requester'), "WH": data.get('wh')
        })
        
    excel_buffer = io.BytesIO()
    df = pd.DataFrame(excel_rows)
    
    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='WRR')
        worksheet = writer.sheets['WRR']
        
        header_fill = PatternFill(start_color='92D050', end_color='92D050', fill_type='solid')
        thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

        for cell in worksheet[1]:
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = thin_border

        for row in worksheet.iter_rows(min_row=2, max_row=worksheet.max_row):
            for cell in row:
                cell.border = thin_border
                col_name = df.columns[cell.column - 1]
                if col_name in ["Item Name [UKR]", "Item Name [ENG]", "WHO Catalogue Item Name"]:
                    cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                else:
                    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=False)

        for i, col in enumerate(df.columns):
            col_letter = get_column_letter(i + 1)
            col_len = max([len(str(x)) for x in df[col].values] + [len(str(col))])
            worksheet.column_dimensions[col_letter].width = min(col_len + 2, 50)

        worksheet.auto_filter.ref = worksheet.dimensions

    excel_buffer.seek(0)
    
    word_buffer = io.BytesIO()
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.line_spacing = 1.15
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.add_run('WORLD HEALTH ORGANIZATION\n').bold = True
    p_title.runs[0].font.size = Pt(16)
    p_title.add_run('RECEIVING REPORT AND ACCEPTANCE INFORMATION').bold = True
    p_title.runs[1].font.size = Pt(14)
    doc.add_paragraph()
    
    header_table = doc.add_table(rows=6, cols=2)
    set_cell_text(header_table.cell(0, 0), "Purchase Order Number:")
    set_cell_text(header_table.cell(0, 1), f"{data.get('po_number', '')}")
    set_cell_text(header_table.cell(1, 0), "Registration Number:")
    set_cell_text(header_table.cell(1, 1), f"{data.get('invoice_info', '')}")
    set_cell_text(header_table.cell(2, 0), "Number of items Received:")
    
    total_qty_str = get_total_quantity(data.get('items', []))
    set_cell_text(header_table.cell(2, 1), total_qty_str)
    
    set_cell_text(header_table.cell(3, 0), "Number of Parcels Received:")
    set_cell_text(header_table.cell(3, 1), f"{data.get('number_of_parcels', '')}")
    set_cell_text(header_table.cell(4, 0), "Item received (check one):")
    set_cell_text(header_table.cell(4, 1), "All [ ] ; Partial [ ]")
    set_cell_text(header_table.cell(5, 0), "Supplier name:")
    set_cell_text(header_table.cell(5, 1), f"{data.get('supplier_name', '')}")
    doc.add_paragraph()
    
    p_decl = doc.add_paragraph()
    p_decl.add_run('DECLARATION').bold = True
    doc.add_paragraph("The consignment against the above-mentioned PO has been received, invoice, packing list and other shipping documents are attached.")
    
    current_date = datetime.datetime.now().strftime("%d.%m.%Y")
    sign_table = doc.add_table(rows=4, cols=2)
    set_cell_text(sign_table.cell(0, 0), "Name of receiver:")
    set_cell_text(sign_table.cell(0, 1), "Oleksandr Tsybulnyk")
    set_cell_text(sign_table.cell(1, 0), "Title:")
    set_cell_text(sign_table.cell(1, 1), "Storekeeper")
    set_cell_text(sign_table.cell(2, 0), "Signature:")
    set_cell_text(sign_table.cell(2, 1), "")
    set_cell_text(sign_table.cell(3, 0), "Date:")
    set_cell_text(sign_table.cell(3, 1), current_date)
    doc.add_paragraph()
    
    p_remarks = doc.add_paragraph()
    p_remarks.add_run('Remarks:').bold = True
    
    if data.get('remark_desc'):
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        set_cell_text(table.rows[0].cells[0], 'Name, dosage form', bold=True)
        set_cell_text(table.rows[0].cells[1], 'Batch#', bold=True)
        set_cell_text(table.rows[0].cells[2], 'Quantity', bold=True)
        set_cell_text(table.rows[0].cells[3], 'Inconsistency description', bold=True)
        table.rows[1].cells[0].text = data.get('remark_item', '')
        table.rows[1].cells[1].text = data.get('remark_batch', '')
        table.rows[1].cells[2].text = data.get('remark_qty', '')
        table.rows[1].cells[3].text = data.get('remark_desc', '')
    else:
        doc.add_paragraph("None")

    doc.save(word_buffer)
    word_buffer.seek(0)
    
    return excel_buffer, word_buffer, base_name

st.title("📦 WHO Warehouse Automation")

tab1, tab2 = st.tabs(["📄 WRR Generator", "📋 Packing List OCR"])

with tab1:
    st.markdown("### 1. Upload WRR Document")
    uploaded_file = st.file_uploader("Upload document scan (PDF, JPG, PNG)", type=['pdf', 'png', 'jpg'], key="wrr_uploader")

    temp_file_path = None
    if uploaded_file is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as temp_file:
            temp_file.write(uploaded_file.read())
            temp_file_path = temp_file.name

        if st.button("🤖 Process Document via AI", use_container_width=True):
            with st.spinner("Extracting data..."):
                st.session_state['extracted_data'] = process_document_with_gemini(temp_file_path)
                
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path)

    if 'extracted_data' in st.session_state and st.session_state['extracted_data']:
        data = st.session_state['extracted_data']
        
        with st.expander("✏️ Review and Edit General Data", expanded=True):
            with st.form("data_form"):
                col1, col2 = st.columns(2)
                with col1:
                    po_number = st.text_input("PO", value=data.get("po_number", ""))
                    act_number = st.text_input("ACT", value=data.get("act_number", ""))
                    or_number = st.text_input("OR", value=data.get("or_number", ""))
                with col2:
                    project = st.text_input("Project", value=data.get("project", ""))
                    task = st.text_input("Task", value=data.get("task", ""))
                    award = st.text_input("Award", value=data.get("award", ""))
                    award_end_date = st.text_input("Award end date", value=data.get("award_end_date", ""))
                    donor = st.text_input("Donor", value=data.get("donor", ""))
                    requester = st.text_input("Requester", value=data.get("requester", ""))
                
                wh = st.text_input("WH (Warehouse)", value=data.get("wh", ""))

                st.markdown("---")
                st.markdown("**WRR Specific Fields**")
                supplier_name = st.text_input("Supplier Name", value=data.get("supplier_name", ""))
                invoice_info = st.text_input("Invoice / Waybill", value=data.get("invoice_info", ""))
                number_of_parcels = st.text_input("Number of Parcels", value=data.get("number_of_parcels", ""))
                
                st.markdown("**Remarks / Discrepancies**")
                col_r1, col_r2 = st.columns(2)
                with col_r1: 
                    remark_item = st.text_input("Remark: Item Name", value="")
                    remark_batch = st.text_input("Remark: Batch", value="")
                with col_r2: 
                    remark_qty = st.text_input("Remark: Quantity", value="")
                    remark_desc = st.text_input("Remark: Inconsistency description", value="")

                st.markdown("---")
                st.markdown("**Items List (Editable Table)**")
                st.info("You can add, delete, or modify any number of rows directly in the table below.")
                
                # Підготовка списку товарів для таблиці
                items_list = data.get("items", [])
                if not items_list:
                    items_list = [{"item_name_ukr": "", "item_name_eng": "", "who_item_code": "", "who_catalogue_item_name": "", "batch": "", "exp_date": "", "quantity": ""}]
                
                df_items = pd.DataFrame(items_list)
                df_items = df_items.rename(columns={
                    "item_name_ukr": "Item Name [UKR]",
                    "item_name_eng": "Item Name [ENG]",
                    "who_item_code": "WHO Item code",
                    "who_catalogue_item_name": "WHO Catalogue Item Name",
                    "batch": "Batch",
                    "exp_date": "Exp.date",
                    "quantity": "Quantity"
                })
                
                cols = ["Item Name [UKR]", "Item Name [ENG]", "WHO Item code", "WHO Catalogue Item Name", "Batch", "Exp.date", "Quantity"]
                for col in cols:
                    if col not in df_items.columns:
                        df_items[col] = ""
                df_items = df_items[cols]
                
                # Інтерактивна таблиця
                edited_items_df = st.data_editor(df_items, num_rows="dynamic", use_container_width=True)

                submitted = st.form_submit_button("✅ Apply Changes & Generate Files", use_container_width=True)

        if submitted:
            final_items = edited_items_df.to_dict('records')
            final_data = {
                "act_number": act_number, "po_number": po_number, "or_number": or_number,
                "project": project, "task": task, "award": award,
                "award_end_date": award_end_date, "donor": donor, "requester": requester, 
                "wh": wh, "supplier_name": supplier_name, "invoice_info": invoice_info,
                "number_of_parcels": number_of_parcels,
                "remark_item": remark_item, "remark_batch": remark_batch, 
                "remark_qty": remark_qty, "remark_desc": remark_desc,
                "items": final_items
            }
            
            excel_buffer, word_buffer, base_name = generate_files_in_memory(final_data)
            
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                zip_file.writestr(f"{base_name}.xlsx", excel_buffer.getvalue())
                zip_file.writestr(f"WRR_{base_name}.docx", word_buffer.getvalue())
            
            # Генерація HTML рядків для всіх товарів
            email_items_html = ""
            for item in final_items:
                email_items_html += f"""<tr>
                <td style="padding: 8px;">{po_number}</td><td style="padding: 8px;">{or_number}</td><td style="padding: 8px;">{requester}</td><td style="padding: 8px;">{item.get('Item Name [ENG]', '')}</td><td style="padding: 8px;">{item.get('Batch', '')}</td><td style="padding: 8px;">{item.get('Exp.date', '')}</td><td style="padding: 8px;">{item.get('Quantity', '')}</td>
                </tr>"""

            st.session_state['wrr_ready'] = True
            st.session_state['wrr_excel_data'] = excel_buffer.getvalue()
            st.session_state['wrr_word_data'] = word_buffer.getvalue()
            st.session_state['wrr_zip_data'] = zip_buffer.getvalue()
            st.session_state['wrr_base_name'] = base_name
            st.session_state['wrr_po_number'] = po_number
            st.session_state['wrr_email_html'] = f"""
            <div style="font-family: Calibri, Arial, sans-serif; font-size: 14px;">
            <p>Dear Team,</p>
            <p>Please find attached the Warehouse Receiving Report (WRR) and Excel database for PO {po_number}.<br>
            Goods have been successfully received at the {wh} warehouse.<br>
            The data has been extracted and transferred to Farmasoft for balance registration.</p>
            <table border="1" style="border-collapse: collapse; width: 100%; max-width: 1000px; border-color: gray;">
            <tr style="background-color: rgba(128, 128, 128, 0.2);">
            <th style="padding: 8px; text-align: left;">PO</th><th style="padding: 8px; text-align: left;">Order request</th><th style="padding: 8px; text-align: left;">Requester</th><th style="padding: 8px; text-align: left;">Item Name [ENG]</th><th style="padding: 8px; text-align: left;">Batch</th><th style="padding: 8px; text-align: left;">Expiry date</th><th style="padding: 8px; text-align: left;">Quantity Inbound</th>
            </tr>{email_items_html}</table></div>
            """

        if st.session_state.get('wrr_ready'):
            st.success("Files successfully generated!")
            st.markdown("### 2. Download Files")
            col_btn1, col_btn2, col_btn3 = st.columns(3) 
            with col_btn1:
                st.download_button(label="📊 Download Excel", data=st.session_state['wrr_excel_data'], file_name=f"{st.session_state['wrr_base_name']}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
            with col_btn2:
                st.download_button(label="📄 Download Word", data=st.session_state['wrr_word_data'], file_name=f"WRR_{st.session_state['wrr_base_name']}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            with col_btn3:
                st.download_button(label="📁 Download Folder (ZIP)", data=st.session_state['wrr_zip_data'], file_name=f"{st.session_state['wrr_base_name']}.zip", mime="application/zip", use_container_width=True)
            
            st.markdown("---")
            st.markdown("### 📧 Email Template")
            st.text_input("Subject:", value=f"Warehouse Receiving Report / PO {st.session_state['wrr_po_number']}")
            st.markdown(st.session_state['wrr_email_html'], unsafe_allow_html=True)
            st.info("💡 Highlight the text and table above, press Ctrl+C (or Copy), and paste it into your email body.")

with tab2:
    st.markdown("### 1. Upload Packing Lists")
    uploaded_pl_files = st.file_uploader("Upload Packing List scans/photos (Select multiple files at once)", type=['pdf', 'png', 'jpg'], accept_multiple_files=True, key="pl_uploader")
    
    if uploaded_pl_files and st.button("🤖 Extract to Excel", use_container_width=True, key="extract_pl"):
        with st.spinner("Processing packing lists via AI..."):
            temp_paths = []
            for f in uploaded_pl_files:
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{f.name.split('.')[-1]}") as temp_file:
                    temp_file.write(f.read())
                    temp_paths.append(temp_file.name)
            
            pl_data = process_packing_lists(temp_paths)
            
            for path in temp_paths:
                if os.path.exists(path):
                    os.remove(path)
            
            if pl_data and "items" in pl_data and len(pl_data["items"]) > 0:
                df_pl = pd.DataFrame(pl_data["items"])
                
                df_pl = df_pl.rename(columns={
                    "carton_no": "Carton no.",
                    "item_no": "Item no.",
                    "quantity": "Quantity",
                    "packing_unit": "Packing unit",
                    "item_description_ukr": "Item description UKR",
                    "item_description_eng": "Item description ENG",
                    "batch_no": "Batch no.",
                    "man_date": "Man. date",
                    "exp_date": "Exp. date"
                })
                
                cols = ["Carton no.", "Item no.", "Quantity", "Packing unit", "Item description UKR", "Item description ENG", "Batch no.", "Man. date", "Exp. date"]
                for col in cols:
                    if col not in df_pl.columns:
                        df_pl[col] = ""
                df_pl = df_pl[cols]
                
                excel_buffer_pl = io.BytesIO()
                with pd.ExcelWriter(excel_buffer_pl, engine='openpyxl') as writer:
                    df_pl.to_excel(writer, index=False, startrow=2, sheet_name='Packing_List')
                    worksheet = writer.sheets['Packing_List']
                    
                    header_fill = PatternFill(start_color='92D050', end_color='92D050', fill_type='solid')
                    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

                    ukr_col_idx = df_pl.columns.get_loc("Item description UKR") + 1
                    module_name = pl_data.get('module_name', '')
                    module_batch = pl_data.get('module_batch', '')
                    batch_text = f"batch: {module_batch}" if module_batch else ""
                    
                    cell_mod = worksheet.cell(row=1, column=ukr_col_idx, value=module_name)
                    cell_batch = worksheet.cell(row=2, column=ukr_col_idx, value=batch_text)
                    cell_mod.alignment = Alignment(horizontal='left', vertical='center')
                    cell_batch.alignment = Alignment(horizontal='left', vertical='center')

                    for cell in worksheet[3]:
                        cell.fill = header_fill
                        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                        cell.border = thin_border

                    for row in worksheet.iter_rows(min_row=4, max_row=worksheet.max_row):
                        for cell in row:
                            cell.border = thin_border
                            col_name = df_pl.columns[cell.column - 1]
                            if col_name in ["Item description UKR", "Item description ENG"]:
                                cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=False)
                            else:
                                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=False)

                    for i, col in enumerate(df_pl.columns):
                        col_letter = get_column_letter(i + 1)
                        col_len = max([len(str(x)) for x in df_pl[col].values] + [len(str(col))])
                        worksheet.column_dimensions[col_letter].width = min(col_len + 2, 50)

                    worksheet.auto_filter.ref = f"A3:{get_column_letter(worksheet.max_column)}{worksheet.max_row}"
                    
                excel_buffer_pl.seek(0)
                
                st.session_state['pl_ready'] = True
                st.session_state['pl_df'] = df_pl
                st.session_state['pl_excel_data'] = excel_buffer_pl.getvalue()
            else:
                st.warning("No items extracted or an error occurred.")

    if st.session_state.get('pl_ready'):
        st.success("Extraction complete! Preview of the data:")
        st.dataframe(st.session_state['pl_df'], use_container_width=True)
        
        st.markdown("### 2. Download Result")
        st.download_button(
            label="📊 Download Formatted Excel Document",
            data=st.session_state['pl_excel_data'],
            file_name="Packing_List.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )